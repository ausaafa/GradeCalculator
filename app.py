import os
import json
import io
from datetime import datetime
from flask import Flask, request, jsonify, render_template, send_file
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

app = Flask(__name__, template_folder="templates")
app.secret_key = os.environ.get("SECRET_KEY", "change-this-in-production")

DEFAULT_EXAM = {
    "Part 1: Pharmaceutical Sciences": {"weight": 25},
    "Part 2: Pharmacy Practice": {"weight": 55},
    "Part 3: Social / Behavioural / Administrative Sciences": {"weight": 20},
}


def safe_num(value, default=0):
    try:
        n = float(value)
        return max(0, min(100, round(n, 1)))
    except (TypeError, ValueError):
        return default


def get_exam_parts(payload):
    """Use the course structure sent from the browser, with a fallback for older clients."""
    course_parts = payload.get("courseParts") or payload.get("examParts") or {}
    if isinstance(course_parts, dict) and course_parts:
        normalized = {}
        for part, meta in course_parts.items():
            if isinstance(meta, dict):
                normalized[part] = {
                    "weight": safe_num(meta.get("weight"), 0),
                    "chapter_count": int(meta.get("chapter_count") or meta.get("chapterCount") or 0),
                }
            else:
                normalized[part] = {"weight": safe_num(meta, 0), "chapter_count": 0}
        return normalized
    return {part: {"weight": meta["weight"], "chapter_count": 0} for part, meta in DEFAULT_EXAM.items()}


@app.route("/")
def home():
    return render_template("index.html")


@app.route("/health")
def health():
    return jsonify({"ok": True})


def compute_basic_results(payload):
    exam_parts = get_exam_parts(payload)
    part_scores = payload.get("partScores", {}) or {}
    chapter_attempts = payload.get("chapterAttempts", {}) or {}
    attempt_timeline = payload.get("attemptTimeline", []) or []

    section_rows = []
    weighted_sum = 0
    coverage = 0

    for part, meta in exam_parts.items():
        weight = safe_num(meta.get("weight"), 0)
        score = part_scores.get(part)

        if score is None:
            chapter_scores = []
            for key, attempts in chapter_attempts.items():
                if not key.startswith(part + "|||"):
                    continue
                if attempts:
                    latest = attempts[-1].get("score")
                    if latest is not None:
                        chapter_scores.append(safe_num(latest))
            score = round(sum(chapter_scores) / len(chapter_scores), 1) if chapter_scores else None

        if score is not None:
            score = safe_num(score)
            contribution = round(score * weight / 100, 1)
            weighted_sum += contribution
            coverage += weight
        else:
            contribution = 0

        section_rows.append({
            "part": part,
            "weight": weight,
            "score": score,
            "contribution": contribution,
            "chapter_count": meta.get("chapter_count", 0),
        })

    weighted_grade = round(weighted_sum / (coverage / 100), 1) if coverage else None
    projected_grade = round(weighted_sum, 1)

    ranked = sorted(
        [row for row in section_rows if row["score"] is not None],
        key=lambda x: x["score"],
        reverse=True,
    )

    vals = [safe_num(x.get("score")) for x in attempt_timeline if x.get("score") is not None]
    avg_attempt = round(sum(vals) / len(vals), 1) if vals else None

    return {
        "weighted_grade": weighted_grade,
        "projected_grade": projected_grade,
        "coverage": round(coverage, 1),
        "avg_attempt": avg_attempt,
        "sections": section_rows,
        "strongest_section": ranked[0]["part"] if ranked else None,
        "weakest_section": ranked[-1]["part"] if ranked else None,
        "total_attempts": len(attempt_timeline),
    }


@app.route("/api/results/basic", methods=["POST"])
def basic_results():
    payload = request.get_json(force=True) or {}
    basic = compute_basic_results(payload)
    return jsonify({"ok": True, "basic": basic})


def build_local_recommendations(basic):
    sections = basic.get("sections", [])
    missing = [s for s in sections if s.get("score") is None]
    weak = sorted([s for s in sections if s.get("score") is not None], key=lambda x: x["score"] or 0)[:2]

    lines = []
    if missing:
        lines.append("Add scores for the missing sections so the final result is based on full coverage.")
    if weak:
        names = ", ".join(s["part"] for s in weak)
        lines.append(f"Prioritize review for: {names}.")
    if basic.get("avg_attempt") is not None:
        if basic["avg_attempt"] < 70:
            lines.append("Redo missed chapter questions and aim for at least 70% before moving on.")
        else:
            lines.append("Keep reviewing missed questions while maintaining your current attempt average.")
    if not lines:
        lines.append("Start entering part scores or chapter attempts to unlock more useful recommendations.")
    return "\n".join(f"- {line}" for line in lines)


def generate_ai_recommendations(payload, basic):
    """Optional OpenAI summary. If OPENAI_API_KEY is not set, the app still works."""
    if not os.environ.get("OPENAI_API_KEY"):
        return None

    try:
        from openai import OpenAI

        client = OpenAI()
        model = os.environ.get("OPENAI_RESULTS_MODEL", "gpt-4.1-mini")
        compact_payload = {
            "course": payload.get("course"),
            "basic_results": basic,
            "recent_attempts": (payload.get("attemptTimeline") or [])[-12:],
        }
        response = client.responses.create(
            model=model,
            input=(
                "Write a short, professional student exam-results summary. "
                "Use plain text only. Include strengths, weak areas, and 3 action steps. "
                "Do not invent scores. Data:\n"
                + json.dumps(compact_payload, ensure_ascii=False)
            ),
        )
        text = getattr(response, "output_text", None)
        return text.strip() if text else None
    except Exception:
        # Never print backend/package errors into student or instructor reports.
        return None


def build_results_text(payload, basic, recommendations):
    generated_at = datetime.now().strftime("%Y-%m-%d %H:%M")
    course = payload.get("course") or "Selected course"
    chapter_attempts = payload.get("chapterAttempts", {}) or {}
    feedback = payload.get("chapterFeedback", {}) or {}

    lines = []
    lines.append("Student Results Summary")
    lines.append("=" * 40)
    lines.append(f"Course: {course}")
    lines.append(f"Generated: {generated_at}")
    lines.append("")
    lines.append("Overall Results")
    lines.append("-" * 16)
    lines.append(f"Current overall grade: {basic['projected_grade']}%")
    lines.append(f"Weighted grade on completed sections: {basic['weighted_grade'] if basic['weighted_grade'] is not None else 'N/A'}")
    lines.append(f"Grade coverage: {basic['coverage']}%")
    lines.append(f"Total attempts: {basic['total_attempts']}")
    lines.append(f"Average attempt score: {basic['avg_attempt'] if basic['avg_attempt'] is not None else 'N/A'}")
    lines.append(f"Strongest section: {basic['strongest_section'] or 'N/A'}")
    lines.append(f"Weakest section: {basic['weakest_section'] or 'N/A'}")
    lines.append("")
    lines.append("Section Breakdown")
    lines.append("-" * 17)
    for row in basic["sections"]:
        score = f"{row['score']}%" if row["score"] is not None else "Not entered"
        lines.append(f"{row['part']}")
        lines.append(f"  Weight: {row['weight']}%")
        lines.append(f"  Score: {score}")
        lines.append(f"  Final-grade contribution: {row['contribution']}%")
    lines.append("")
    lines.append("Latest Chapter Attempts")
    lines.append("-" * 23)

    latest_rows = []
    for key, attempts in chapter_attempts.items():
        if not attempts:
            continue
        latest = attempts[-1]
        part, chapter = key.split("|||", 1) if "|||" in key else (latest.get("part", ""), latest.get("chapter", key))
        latest_rows.append((part, chapter, latest))

    if latest_rows:
        for part, chapter, latest in latest_rows:
            lines.append(f"{chapter}")
            lines.append(f"  Section: {part}")
            lines.append(f"  Latest score: {latest.get('earned')}/{latest.get('possible')} ({latest.get('score')}%)")
            if latest.get("quickSource"):
                lines.append(f"  Source: {latest.get('quickSource')}")
    else:
        lines.append("No chapter attempts entered yet.")

    if feedback:
        reaction_labels = {"up": "above 70%", "neutral": "50-60%", "down": "below 50%"}
        lines.append("")
        lines.append("Emoji Quick-Score Notes")
        lines.append("-" * 23)
        for key, reaction in feedback.items():
            chapter = key.split("|||", 1)[1] if "|||" in key else key
            lines.append(f"{chapter}: {reaction_labels.get(reaction, reaction)}")

    if recommendations:
        lines.append("")
        lines.append("Recommendations")
        lines.append("-" * 15)
        lines.append(recommendations)
    lines.append("")
    return "\n".join(lines)



def fmt_percent(value):
    if value is None:
        return "N/A"
    return f"{safe_num(value)}%"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table_header(row):
    for cell in row.cells:
        set_cell_shading(cell, "111827")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)


def add_table(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(header_cells[i], header, bold=True)
    style_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    return table


def latest_chapter_rows(payload):
    chapter_attempts = payload.get("chapterAttempts", {}) or {}
    rows = []
    for key, attempts in chapter_attempts.items():
        if not attempts:
            continue
        latest = attempts[-1]
        part, chapter = key.split("|||", 1) if "|||" in key else (latest.get("part", ""), latest.get("chapter", key))
        rows.append({
            "part": part,
            "chapter": chapter,
            "earned": latest.get("earned"),
            "possible": latest.get("possible"),
            "score": safe_num(latest.get("score")),
            "quick_source": latest.get("quickSource"),
            "timestamp": latest.get("timestamp", ""),
        })
    return rows


def analyze_results(payload, basic):
    sections = basic.get("sections", [])
    scored = [s for s in sections if s.get("score") is not None]
    strengths = sorted([s for s in scored if safe_num(s.get("score")) >= 70], key=lambda x: x["score"], reverse=True)
    weaknesses = sorted([s for s in scored if safe_num(s.get("score")) < 70], key=lambda x: x["score"])
    missing = [s for s in sections if s.get("score") is None]

    timeline = payload.get("attemptTimeline", []) or []
    trend_scores = [safe_num(item.get("score")) for item in timeline if item.get("score") is not None]
    trend = "Not enough chapter attempts to calculate a trend yet."
    if len(trend_scores) >= 2:
        delta = round(trend_scores[-1] - trend_scores[0], 1)
        if delta > 5:
            direction = "improving"
        elif delta < -5:
            direction = "declining"
        else:
            direction = "stable"
        trend = f"The chapter attempt trend is {direction}. First attempt: {trend_scores[0]}%, latest attempt: {trend_scores[-1]}%, change: {delta:+.1f}%."

    chapter_rows = latest_chapter_rows(payload)
    weak_chapters = sorted([r for r in chapter_rows if r["score"] < 70], key=lambda x: x["score"])[:8]
    strong_chapters = sorted([r for r in chapter_rows if r["score"] >= 70], key=lambda x: x["score"], reverse=True)[:8]

    return {
        "strengths": strengths,
        "weaknesses": weaknesses,
        "missing": missing,
        "trend": trend,
        "weak_chapters": weak_chapters,
        "strong_chapters": strong_chapters,
    }


def make_results_charts(payload, basic):
    """Return chart image buffers. If matplotlib is unavailable, report still generates."""
    charts = {}
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
    except Exception:
        return charts

    sections = basic.get("sections", [])
    labels = [s["part"].split(":")[0][:18] for s in sections]
    scores = [safe_num(s.get("score")) if s.get("score") is not None else 0 for s in sections]
    weights = [safe_num(s.get("weight")) for s in sections]

    if sections:
        fig, ax = plt.subplots(figsize=(7.2, 3.2), dpi=150)
        ax.bar(labels, weights, label="Section weight")
        ax.bar(labels, [round(weights[i] * scores[i] / 100, 1) for i in range(len(scores))], label="Earned points")
        ax.set_title("Section Weight vs Earned Grade Points")
        ax.set_ylabel("Final-grade weight (%)")
        ax.legend(fontsize=8)
        ax.tick_params(axis="x", labelrotation=15)
        fig.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format="png", bbox_inches="tight")
        plt.close(fig)
        buf.seek(0)
        charts["sections"] = buf

    timeline = payload.get("attemptTimeline", []) or []
    trend_scores = [safe_num(item.get("score")) for item in timeline if item.get("score") is not None]
    if trend_scores:
        fig, ax = plt.subplots(figsize=(7.2, 3.0), dpi=150)
        ax.plot(range(1, len(trend_scores) + 1), trend_scores, marker="o", linewidth=2, label="Attempt score")
        ax.axhline(70, linestyle="--", linewidth=1, label="70% target")
        ax.set_title("Chapter Attempt Trend")
        ax.set_xlabel("Attempt order")
        ax.set_ylabel("Score (%)")
        ax.set_ylim(0, 100)
        ax.legend(fontsize=8)
        fig.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format="png", bbox_inches="tight")
        plt.close(fig)
        buf.seek(0)
        charts["trend"] = buf

    return charts


def build_results_docx(payload, basic, recommendations):
    course = payload.get("course") or "Selected course"
    generated_at = datetime.now().strftime("%Y-%m-%d %H:%M")
    analysis = analyze_results(payload, basic)

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9.5)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Instructor Results Report")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(17, 24, 39)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"Course: {course}  |  Generated: {generated_at}").font.size = Pt(10)

    document.add_paragraph()
    document.add_heading("Executive Summary", level=1)
    summary_rows = [
        ["Current overall grade", fmt_percent(basic.get("projected_grade"))],
        ["Weighted grade on completed sections", fmt_percent(basic.get("weighted_grade"))],
        ["Grade coverage", fmt_percent(basic.get("coverage"))],
        ["Total attempts", basic.get("total_attempts", 0)],
        ["Average attempt score", fmt_percent(basic.get("avg_attempt"))],
        ["Strongest section", basic.get("strongest_section") or "N/A"],
        ["Weakest section", basic.get("weakest_section") or "N/A"],
    ]
    add_table(document, ["Metric", "Result"], summary_rows)

    document.add_paragraph()
    document.add_heading("Trend Insight", level=1)
    document.add_paragraph(analysis["trend"])

    charts = make_results_charts(payload, basic)
    if charts.get("sections"):
        document.add_picture(charts["sections"], width=Inches(6.8))
    if charts.get("trend"):
        document.add_picture(charts["trend"], width=Inches(6.8))

    document.add_heading("Section Breakdown", level=1)
    section_rows = []
    for row in basic.get("sections", []):
        score = fmt_percent(row.get("score")) if row.get("score") is not None else "Not entered"
        section_rows.append([
            row.get("part", ""),
            fmt_percent(row.get("weight")),
            score,
            fmt_percent(row.get("contribution")),
            row.get("chapter_count", 0),
        ])
    add_table(document, ["Section", "Weight", "Score", "Contribution", "Chapters"], section_rows)

    document.add_paragraph()
    document.add_heading("Strengths", level=1)
    if analysis["strengths"]:
        for item in analysis["strengths"][:5]:
            document.add_paragraph(f"{item['part']}: {fmt_percent(item['score'])}. Maintain this area with quick spaced review and missed-question checks.", style="List Bullet")
    else:
        document.add_paragraph("No section is above the 70% target yet. Focus on building consistency across the lowest-weight and highest-weight weak areas first.")

    if analysis["strong_chapters"]:
        document.add_paragraph("Strong chapter signals:")
        for row in analysis["strong_chapters"]:
            document.add_paragraph(f"{row['chapter']} — {fmt_percent(row['score'])}", style="List Bullet")

    document.add_heading("Weaknesses and Instructor Focus Areas", level=1)
    if analysis["weaknesses"]:
        for item in analysis["weaknesses"][:6]:
            document.add_paragraph(f"{item['part']}: {fmt_percent(item['score'])}. Prioritize targeted review, extra practice questions, and re-testing.", style="List Bullet")
    else:
        document.add_paragraph("No scored section is currently below 70%. Continue monitoring new attempts for dips.")

    if analysis["weak_chapters"]:
        document.add_paragraph("Weak chapter signals:")
        weak_rows = [[r["chapter"], r["part"], fmt_percent(r["score"]), r.get("quick_source") or "Manual score"] for r in analysis["weak_chapters"]]
        add_table(document, ["Chapter", "Section", "Latest Score", "Source"], weak_rows)

    if analysis["missing"]:
        document.add_heading("Missing Data", level=1)
        for item in analysis["missing"]:
            document.add_paragraph(f"{item['part']}: no part score or chapter aggregate yet. Add data before making a final readiness decision.", style="List Bullet")

    if recommendations:
        document.add_heading("Instructor Summary and Next Steps", level=1)
        for line in str(recommendations).splitlines():
            clean = line.strip().lstrip("- ").strip()
            if clean:
                document.add_paragraph(clean, style="List Bullet")

    document.add_paragraph()
    footer = document.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Generated by Exam Assessment Evaluator")
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 116, 139)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


@app.route("/api/results/download", methods=["POST"])
def download_results():
    payload = request.get_json(force=True) or {}
    basic = compute_basic_results(payload)

    # Default behavior: do not call OpenAI for the downloadable report.
    # This prevents messages like "No module named 'openai'" from appearing in DOCX files.
    # If you ever want AI text later, send {"includeAiRecommendations": true}.
    recommendations = None
    include_ai = bool(payload.get("includeAiRecommendations") or payload.get("includeAIRecommendations"))
    include_recommendations = payload.get("includeRecommendations", True) is not False

    if include_recommendations:
        if include_ai:
            recommendations = generate_ai_recommendations(payload, basic)
        recommendations = recommendations or build_local_recommendations(basic)

    report = build_results_docx(payload, basic, recommendations)

    safe_course = "student-results"
    if payload.get("course"):
        safe_course = "".join(ch.lower() if ch.isalnum() else "-" for ch in payload["course"]).strip("-")[:70]
    filename = f"{safe_course}-results-report.docx"

    return send_file(
        report,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=filename,
    )


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5055))
    app.run(host="0.0.0.0", port=port, debug=True)
